import os
from openai import OpenAI
import streamlit as st
import pandas as pd
import time
import openpyxl
import plotly.express as px
import plotly.graph_objects as go
from docx import Document


# Initialize the OpenAI client with the API key from Streamlit secrets
api_key = st.secrets["OPENAI_API_KEY"]
client = OpenAI(api_key=api_key)

def createAssistant(file_ids, title):
    instructions = """You are an assistant designed to answer questions based on the content of the uploaded PDF and Excel files. Follow these guidelines to provide accurate and relevant responses:

        1. **Content Extraction:**
        - Extract information directly from the text, data, and figures within the uploaded PDF files.
        - Extract data, tables, and charts from Excel files, interpreting them as needed.
        - Interpret and summarize sections of the PDFs and Excel spreadsheets to answer questions.
        - While responding to a question, you must also give a reference from which file you have taken the response and then mention it.

        2. **Answering Questions:**
        - Provide clear and concise answers to the user's questions using the content from the PDFs and Excel files.
        - Ensure that your responses are relevant and directly supported by the information in the files.

        3. **Citation and Referencing:**
        - When providing answers, include citations that reference specific sections, pages, or figures from the PDFs where the information was found.
        - For Excel files, refer to specific sheets, cells, or ranges that were used in forming the response.
        - Format citations clearly, e.g., "As shown in Document X, page Y..." or "According to Sheet Z, cell A1..."

        4. **Handling Multiple Files:**
        - If multiple PDF or Excel files are uploaded, consider all available files when answering questions.
        - Indicate which file and section or sheet the information comes from if relevant.

        5. **Accuracy and Clarity:**
        - Only provide information that is present in the uploaded PDFs or Excel files.
        - Avoid speculation or fabricated details.
        - Use bullet points or lists to enhance readability if needed.

        6. **If Information is Missing:**
        - If the information needed to answer a question is not found in the PDFs or Excel files, clearly state that the information is not available in the provided files.

        7. **User Interaction:**
        - Maintain a polite and professional tone.
        - Encourage users to upload additional files if the provided PDFs or Excel spreadsheets do not contain the necessary information.

        8. **References at the End of Responses:**
        - At the end of every response, list the reference(s) indicating which file(s) the information was extracted from, using a format like: "References: Document X, page Y" or "References: Spreadsheet Z, sheet S, cell range A1:B2."

        Follow these instructions to ensure that you provide helpful and accurate responses based on the content of the uploaded PDFs and Excel files."""

    model = "gpt-4o-mini"
    tools = [{"type": "file_search"}]
    vector_store = client.beta.vector_stores.create(name=title, file_ids=file_ids)
    tool_resources = {"file_search": {"vector_store_ids": [vector_store.id]}}

    assistant = client.beta.assistants.create(
        name=title,
        instructions=instructions,
        model=model,
        tools=tools,
        tool_resources=tool_resources
    )

    return assistant.id, vector_store.id

def saveFileOpenAI(location):
    with open(location, "rb") as file:
        uploaded_file = client.files.create(file=file, purpose='assistants')
    return uploaded_file.id

def startAssistantThread(prompt, vector_id):
    messages = [{"role": "user", "content": prompt}]
    tool_resources = {"file_search": {"vector_store_ids": [vector_id]}}
    thread = client.beta.threads.create(messages=messages, tool_resources=tool_resources)
    return thread.id

def runAssistant(thread_id, assistant_id):
    run = client.beta.threads.runs.create(thread_id=thread_id, assistant_id=assistant_id)
    return run.id

def checkRunStatus(thread_id, run_id):
    run = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run_id)
    return run.status

def retrieveThread(thread_id):
    thread_messages = client.beta.threads.messages.list(thread_id)
    list_messages = thread_messages.data
    thread_messages = []
    for message in list_messages:
        obj = {}
        obj['content'] = message.content[0].text.value
        obj['role'] = message.role
        thread_messages.append(obj)
    return thread_messages[::-1]

def addMessageToThread(thread_id, prompt):
    client.beta.threads.messages.create(thread_id, role="user", content=prompt)


def process_run(thread_id, assistant_id):
    run_id = runAssistant(thread_id, assistant_id)
    status = 'running'

    while status != 'completed':
        time.sleep(2)
        status = checkRunStatus(thread_id, run_id)

    thread_messages = retrieveThread(thread_id)
    return thread_messages[-1]  # Return only the last message

def start_new_chat():
    st.session_state.chat_history = []
    if 'thread_id' in st.session_state:
        del st.session_state.thread_id
    st.session_state.user_input = ""

def handle_input(user_message):
    if user_message:
        if "thread_id" not in st.session_state:
            if "vector_id" not in st.session_state:
                st.error("Please upload a file first before starting a chat.")
                return
            # Start a new assistant thread
            st.session_state.thread_id = startAssistantThread(user_message, st.session_state.vector_id)
        else:
            # Add message to existing thread
            addMessageToThread(st.session_state.thread_id, user_message)

        # Append the user message to the chat history
        st.session_state.chat_history.append({"role": "user", "content": user_message})

        # Process the assistant's response
        assistant_response = process_run(
            st.session_state.thread_id, st.session_state.assistant_id
        )

        # Append the assistant response to the chat history
        st.session_state.chat_history.append(assistant_response)

        # Force a rerun to update the chat display
        st.rerun()

def toggle_theme():
    st.session_state.theme = 'dark' if st.session_state.theme == 'light' else 'light'
    st.rerun()

def convert_excel_to_docx(excel_file):
    wb = openpyxl.load_workbook(excel_file)
    doc = Document()
    
    for sheet in wb.worksheets:
        doc.add_heading(sheet.title, level=1)
        for row in sheet.iter_rows(values_only=True):
            doc.add_paragraph('\t'.join(str(cell) for cell in row))
    
    location = f"temp_file_{excel_file.name}.docx"
    doc.save(location)
    return location

def extract_data_from_excel(uploaded_file):
    # Read Excel file into a DataFrame
    df = pd.read_excel(uploaded_file , header = 0)
    return df



def create_graph():
    if 'excel_data' not in st.session_state:
        st.warning("Please upload an Excel file first.")
        return

    df = st.session_state.excel_data
    col_names = df.columns.tolist()

    # Ensure that all numeric columns are actually numeric
    df_numeric = df.copy()
    for col in df_numeric.select_dtypes(include=['object']).columns:
        df_numeric[col] = pd.to_numeric(df_numeric[col], errors='coerce')

    x_axis = st.selectbox("Select X-axis column", col_names)
    y_axis = st.selectbox("Select Y-axis column", col_names)

    chart_types = [
        "Line", "Bar", "Scatter", "Pie", 
        "Area", "Histogram", "Box", "Bubble", 
        "Radar", "Gantt"
    ]
    chart_type = st.selectbox("Select chart type", chart_types)

    # Generate chart based on the selected type and options
    if chart_type == "Line":
        line_mode = st.selectbox("Select line mode", ["lines", "markers", "lines+markers"])
        line_shape = st.selectbox("Select line shape", ["linear", "spline", "vhv", "hvh", "vh", "hv"])
        fig = px.line(df_numeric, x=x_axis, y=y_axis, line_shape=line_shape)
        fig.update_traces(mode=line_mode)

    elif chart_type == "Bar":
        barmode = st.selectbox("Select bar mode", ["group", "stack", "overlay", "relative"])
        orientation = st.selectbox("Select orientation", ["vertical", "horizontal"])
        fig = px.bar(df_numeric, x=x_axis, y=y_axis, orientation='h' if orientation == "horizontal" else 'v')
        fig.update_layout(barmode=barmode)

    elif chart_type == "Scatter":
        marker_size = st.slider("Select marker size", 1, 20, 5)
        marker_symbol = st.selectbox("Select marker symbol", px.symbols)
        fig = px.scatter(df_numeric, x=x_axis, y=y_axis, size_max=marker_size, symbol=marker_symbol)

    elif chart_type == "Pie":
        hole_size = st.slider("Select hole size (for donut chart effect)", 0.0, 0.9, 0.0)
        fig = px.pie(df, names=x_axis, values=y_axis, hole=hole_size)

    elif chart_type == "Area":
        line_shape = st.selectbox("Select line shape", ["linear", "spline"])
        fig = px.area(df_numeric, x=x_axis, y=y_axis, line_shape=line_shape)

    elif chart_type == "Histogram":
        nbins = st.slider("Select number of bins", 5, 50, 20)
        hist_norm = st.selectbox("Select normalization", ["", "percent", "probability", "density", "probability density"])
        fig = px.histogram(df_numeric, x=x_axis, nbins=nbins, histnorm=hist_norm)

    elif chart_type == "Box":
        box_mode = st.selectbox("Select box mode", ["group", "overlay"])
        fig = px.box(df_numeric, x=x_axis, y=y_axis)
        fig.update_traces(boxmode=box_mode)

    elif chart_type == "Bubble":
        size = st.selectbox("Select size column", col_names)
        fig = px.scatter(df_numeric, x=x_axis, y=y_axis, size=size, color=x_axis)

    elif chart_type == "Radar":
        fig = go.Figure()
        for index, row in df.iterrows():
            fig.add_trace(go.Scatterpolar(r=row.values, theta=row.index, fill='toself', name=row.name))
        fig.update_layout(polar=dict(radialaxis=dict(visible=True)))

    elif chart_type == "Gantt":
        if 'Start' in df.columns and 'Finish' in df.columns:
            fig = px.timeline(df, x_start='Start', x_end='Finish', y=x_axis, title='Gantt Chart')
            fig.update_yaxes(categoryorder="total ascending")
        else:
            st.warning("For Gantt chart, ensure 'Start' and 'Finish' columns exist in the data.")
            return

    st.plotly_chart(fig)
